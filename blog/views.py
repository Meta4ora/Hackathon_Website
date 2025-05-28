# Импорт необходимых библиотек и модулей для работы с Django, базами данных, отчетами и логированием
import json
import os
from django.conf import settings
from datetime import datetime
from django.shortcuts import render, redirect, get_object_or_404
from django.core.exceptions import ValidationError
from django.db import connections, IntegrityError, DatabaseError, connection
from django.views.decorators.http import require_http_methods, require_GET
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Импорт форм и моделей, используемых в представлениях
from .forms import LoginForm, RegisterForm, ProfileForm, FeedbackForm, TeamRegistrationForm, MentorRegistrationForm
from .models import PublicEvent
from django.contrib import messages
from django.utils.timezone import now
import io
import logging
from django.http import JsonResponse, FileResponse
from django.views.decorators.http import require_POST
from django.db import connections, DatabaseError
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
import docx
import time

# Представление главной страницы
def index(request):
    # Получение последних 6 публичных мероприятий
    events = PublicEvent.objects.all().order_by('-start_date')[:6]

    # Проверка аутентификации пользователя
    is_authenticated = 'user_role' in request.session and 'id' in request.session
    user_data = {
        'email': request.session.get('email'),
        'role': request.session.get('user_role'),
    } if is_authenticated else None

    # Получение предстоящих мероприятий без ментора
    with connection.cursor() as cursor:
        cursor.execute("""
            SELECT id_event, event_name, start_date
            FROM events
            WHERE start_date > CURRENT_DATE AND id_mentor IS NULL
            ORDER BY start_date ASC
        """)
        upcoming_events = [
            {"id_event": row[0], "event_name": row[1], "start_date": row[2]}
            for row in cursor.fetchall()
        ]

    # Получение последних 20 отзывов
    with connection.cursor() as cursor:
        cursor.execute("""
            SELECT fb.feedback_text, fb.feedback_date, ev.event_name,
                   p.name, p.surname
            FROM feedback fb
            JOIN participants p ON fb.id_participant = p.id_participant
            JOIN eventdetailsview ev ON fb.id_event = ev.id_event
            ORDER BY fb.feedback_date DESC
            LIMIT 20
        """)
        feedback_list = [
            {
                'text': row[0],
                'date': row[1],
                'event_name': row[2],
                'first_name': row[3],
                'last_name': row[4],
            } for row in cursor.fetchall()
        ]

    # Получение случайных призеров для каждого места (1, 2, 3)
    winners = []
    with connections['guest'].cursor() as cursor:
        for place in [1, 2, 3]:
            cursor.execute("""
                SELECT c.place_achieved, c.monetary_reward, c.event_name,
                       p.surname, p.name, p.patronymic
                FROM public.certificates c
                JOIN public.participants p ON c.id_participant = p.id_participant
                WHERE c.place_achieved = %s
                ORDER BY RANDOM()
                LIMIT 1
            """, [place])
            row = cursor.fetchone()
            if row:
                winners.append({
                    'place': row[0],
                    'award': row[1],
                    'event_name': row[2],
                    'last_name': row[3],
                    'first_name': row[4],
                    'patronymic': row[5] or '',
                })

    # Проверка необходимости отображения модального окна для регистрации ментора
    show_mentor_modal = request.session.pop('show_mentor_modal', False)

    # Рендеринг шаблона главной страницы с передачей данных
    return render(request, 'index.html', {
        'events': events,
        'is_authenticated': is_authenticated,
        'user_data': user_data,
        'upcoming_events': upcoming_events,
        'feedback_list': feedback_list,
        'winners': winners,
        'now': now(),
        'show_mentor_modal': show_mentor_modal
    })

# Представление для регистрации нового пользователя
def register(request):
    if request.method == 'POST':
        # Создание формы с данными из POST-запроса
        form = RegisterForm(request.POST)
        if form.is_valid():
            # Сохранение данных пользователя и установка сессии
            user_data = form.save()
            request.session['user_role'] = user_data['role']
            request.session['id'] = user_data['id_participant']
            request.session['email'] = user_data['email']
            messages.success(request, "Регистрация прошла успешно! Вы вошли в систему.")
            return redirect('index')
    else:
        # Создание пустой формы для GET-запроса
        form = RegisterForm()
    # Рендеринг шаблона регистрации
    return render(request, 'register.html', {'form': form})

# Представление для аутентификации пользователя
def custom_login(request):
    if request.method == 'POST':
        # Создание формы с данными из POST-запроса
        form = LoginForm(request.POST)
        if form.is_valid():
            try:
                # Аутентификация пользователя
                role, user_data = form.authenticate_user()
                request.session['user_role'] = role
                # Установка ID в зависимости от роли
                if role == 'admin':
                    request.session['id'] = user_data['id_staff']
                else:
                    request.session['id'] = user_data['id_participant']
                request.session['email'] = user_data['email']
                return redirect('index')
            except ValidationError as e:
                # Добавление ошибки в форму при неверных данных
                form.add_error(None, e)
    else:
        # Создание пустой формы для GET-запроса
        form = LoginForm()
    # Рендеринг шаблона авторизации
    return render(request, 'authorization.html', {'form': form})

# Представление для выхода из системы
def custom_logout(request):
    # Очистка сессии
    request.session.flush()
    return redirect('index')

# Представление для отображения списка мероприятий
def events(request):
    # Получение всех мероприятий, отсортированных по дате начала
    events = PublicEvent.objects.all().order_by('-start_date')
    # Проверка аутентификации
    is_authenticated = 'user_role' in request.session and 'id' in request.session
    user_data = {
        'email': request.session.get('email'),
        'role': request.session.get('user_role'),
    } if is_authenticated else None
    # Рендеринг шаблона мероприятий
    return render(request, 'events.html', {
        'events': events,
        'is_authenticated': is_authenticated,
        'user_data': user_data,
        'now': now()
    })

# Представление для отображения деталей мероприятия
def event_detail(request, event_id):
    # Получение мероприятия по ID или возврат 404
    event = get_object_or_404(PublicEvent, id_event=event_id)
    # Проверка аутентификации
    is_authenticated = 'user_role' in request.session and 'id' in request.session
    user_data = {
        'email': request.session.get('email'),
        'role': request.session.get('user_role'),
    } if is_authenticated else None
    # Рендеринг шаблона деталей мероприятия
    return render(request, 'event_detail.html', {
        'event': event,
        'is_authenticated': is_authenticated,
        'user_data': user_data,
        'now': now()
    })

# Представление для добавления нового мероприятия (только POST)
@require_http_methods(["POST"])
def add_event(request):
    # Получение данных из POST-запроса
    event_name = request.POST.get('event_name')
    start_date = request.POST.get('start_date')
    end_date = request.POST.get('end_date')
    id_venue = request.POST.get('id_venue')

    try:
        with connections['default'].cursor() as cursor:
            # Получение максимального ID мероприятия
            cursor.execute("SELECT MAX(id_event) FROM public.events")
            max_id = cursor.fetchone()[0] or 0
            new_id = max_id + 1

            # Вставка нового мероприятия с id_mentor = NULL
            cursor.execute("""
                INSERT INTO public.events (id_event, event_name, start_date, end_date, id_venue, id_mentor)
                VALUES (%s, %s, %s, %s, %s, NULL)
            """, [new_id, event_name, start_date, end_date, id_venue])

        messages.success(request, f"Мероприятие успешно добавлено с ID {new_id}.")

    except (IntegrityError, DatabaseError) as e:
        # Обработка ошибок базы данных
        messages.error(request, f"Ошибка при добавлении мероприятия: {str(e)}")

    return redirect('events')

# Инициализация логгера для отладки
logger = logging.getLogger(__name__)

# Представление для страницы профиля пользователя
def profile(request):
    # Проверка аутентификации
    if 'user_role' not in request.session or 'id' not in request.session:
        messages.error(request, "Пожалуйста, войдите в систему.")
        return redirect('login')

    # Формирование данных пользователя из сессии
    user_data = {
        'email': request.session.get('email'),
        'user_role': request.session.get('user_role'),
        'id': request.session.get('id'),
    }

    # Получение данных пользователя из базы данных
    with connections['org_db'].cursor() as cursor:
        if user_data['user_role'] == 'admin':
            cursor.execute("SELECT * FROM public.staff WHERE id_staff = %s", [user_data['id']])
        else:
            cursor.execute("SELECT * FROM public.participants WHERE id_participant = %s", [user_data['id']])
        row = cursor.fetchone()
        if not row:
            logger.error(f"No user found for id: {user_data['id']} with role: {user_data['user_role']}")
            messages.error(request, "Пользователь не найден.")
            return redirect('custom_login')

        # Формирование словаря с данными пользователя
        columns = [col[0] for col in cursor.description]
        user_details = dict(zip(columns, row))
        user_data.update(user_details)
        user_data['role'] = user_data['user_role']
        logger.debug(f"user_data after update: {user_data}")

    # Очистка номера телефона от пробелов
    if 'phone' in user_data:
        user_data['phone'] = user_data['phone'].strip()

    # Обработка даты рождения
    birth_date = user_data.get('birth_date')
    if isinstance(birth_date, str):
        try:
            birth_date = datetime.strptime(birth_date, '%Y-%m-%d').date()
        except ValueError:
            logger.warning(f"Invalid birth_date format for user {user_data['id']}: {birth_date}")
            birth_date = None
    user_data['birth_date'] = birth_date

    logger.debug(f"user_data: {user_data}")

    # Инициализация форм
    initial_data = {
        'name': user_data.get('name', ''),
        'surname': user_data.get('surname', ''),
        'patronymic': user_data.get('patronymic', ''),
        'birth_date': user_data.get('birth_date').strftime('%Y-%m-%d') if user_data.get('birth_date') else None,
        'phone': user_data.get('phone', ''),
        'role': user_data.get('role', ''),
    }
    profile_form = ProfileForm(initial=initial_data)
    feedback_form = FeedbackForm(id_participant=user_data['id'] if user_data['user_role'] == 'participant' else None)
    team_form = TeamRegistrationForm()
    logger.debug(f"Feedback form initialized for id_participant: {user_data['id'] if user_data['user_role'] == 'participant' else None}")

    if request.method == 'POST':
        if 'profile_submit' in request.POST:
            # Обработка формы профиля
            profile_form = ProfileForm(request.POST)
            if profile_form.is_valid():
                try:
                    cleaned_data = profile_form.cleaned_data
                    with connections['org_db'].cursor() as cursor:
                        if user_data['user_role'] == 'admin':
                            # Обновление данных администратора
                            cursor.execute("""
                                UPDATE public.staff
                                SET name = %s, surname = %s, patronymic = %s, birth_date = %s, phone = %s, role = %s
                                WHERE id_staff = %s
                            """, [
                                cleaned_data['name'],
                                cleaned_data['surname'],
                                cleaned_data['patronymic'],
                                cleaned_data['birth_date'],
                                cleaned_data['phone'],
                                cleaned_data['role'],
                                user_data['id']
                            ])
                        else:
                            # Обновление данных участника
                            cursor.execute("""
                                UPDATE public.participants
                                SET name = %s, surname = %s, patronymic = %s, birth_date = %s, phone = %s, role = %s
                                WHERE id_participant = %s
                            """, [
                                cleaned_data['name'],
                                cleaned_data['surname'],
                                cleaned_data['patronymic'],
                                cleaned_data['birth_date'],
                                cleaned_data['phone'],
                                cleaned_data['role'],
                                user_data['id']
                            ])
                    messages.success(request, "Данные профиля успешно обновлены!")
                    return redirect('profile')
                except (IntegrityError, DatabaseError) as e:
                    logger.error(f"Database error during profile update: {str(e)}")
                    messages.error(request, f"Ошибка при обновлении данных: {str(e)}")
            else:
                logger.debug(f"Profile form errors: {profile_form.errors}")
                messages.error(request, "Пожалуйста, исправьте ошибки в форме профиля.")

        elif 'feedback_submit' in request.POST and user_data['user_role'] == 'participant':
            # Обработка формы отзыва
            feedback_form = FeedbackForm(request.POST, id_participant=user_data['id'])
            if feedback_form.is_valid():
                try:
                    cleaned_data = feedback_form.cleaned_data
                    with connections['org_db'].cursor() as cursor:
                        # Получение следующего ID для отзыва
                        cursor.execute("SELECT MAX(id_feedback) FROM public.feedback")
                        max_id = cursor.fetchone()[0] or 0
                        new_id = max_id + 1

                        # Вставка отзыва в базу данных
                        cursor.execute("""
                            INSERT INTO public.feedback (id_feedback, id_participant, id_event, feedback_text, feedback_date)
                            VALUES (%s, %s, %s, %s, %s)
                        """, [
                            new_id,
                            user_data['id'],
                            cleaned_data['event'],
                            cleaned_data['feedback_text'],
                            now()
                        ])
                    messages.success(request, "Отзыв успешно добавлен!")
                    request.session['show_feedback_modal'] = True
                    return redirect('profile')
                except (IntegrityError, DatabaseError) as e:
                    logger.error(f"Database error during feedback submission: {str(e)}")
                    messages.error(request, f"Ошибка при добавлении отзыва: {str(e)}")
            else:
                logger.debug(f"Feedback form errors: {feedback_form.errors}")
                messages.error(request, "Пожалуйста, исправьте ошибки в форме отзыва.")

        elif 'team_submit' in request.POST and user_data['user_role'] == 'participant':
            # Обработка формы регистрации команды
            team_form = TeamRegistrationForm(request.POST)
            if team_form.is_valid():
                try:
                    cleaned_data = team_form.cleaned_data
                    with connections['org_db'].cursor() as cursor:
                        # Получение следующего ID для команды
                        cursor.execute("SELECT MAX(id_team) FROM public.teams")
                        max_id = cursor.fetchone()[0] or 0
                        new_team_id = max_id + 1

                        # Вставка новой команды
                        captain_name = f"{user_data['surname']} {user_data['name']} {user_data['patronymic']}"
                        cursor.execute("""
                            INSERT INTO public.teams (id_team, status, team_name, captain)
                            VALUES (%s, %s, %s, %s)
                        """, [
                            new_team_id,
                            cleaned_data['status'],
                            cleaned_data['team_name'],
                            captain_name
                        ])

                        # Связывание команды с мероприятием
                        cursor.execute("""
                            INSERT INTO public.event_teams (id_team, id_event)
                            VALUES (%s, %s)
                        """, [
                            new_team_id,
                            cleaned_data['event']
                        ])

                        # Связывание участников с командой
                        participant_ids = [p['id_participant'] for p in cleaned_data['participants']]
                        for participant_id in participant_ids:
                            cursor.execute("""
                                INSERT INTO public.team_participant (id_team, id_participant)
                                VALUES (%s, %s)
                            """, [
                                new_team_id,
                                participant_id
                            ])

                    messages.success(request, "Команда успешно зарегистрирована!")
                    request.session['show_team_modal'] = True
                    return redirect('profile')
                except (IntegrityError, DatabaseError) as e:
                    logger.error(f"Database error during team registration: {str(e)}")
                    messages.error(request, f"Ошибка при регистрации команды: {str(e)}")
            else:
                logger.debug(f"Team form errors: {team_form.errors}")
                messages.error(request, "Пожалуйста, исправьте ошибки в форме регистрации команды.")

    # Очистка флагов модальных окон
    show_feedback_modal = request.session.pop('show_feedback_modal', False)
    show_team_modal = request.session.pop('show_team_modal', False)

    # Рендеринг шаблона профиля
    return render(request, 'profile.html', {
        'profile_form': profile_form,
        'feedback_form': feedback_form,
        'team_form': team_form,
        'user_data': user_data,
        'show_feedback_modal': show_feedback_modal,
        'show_team_modal': show_team_modal
    })

# Представление для панели управления (только для администраторов)
def control_panel(request):
    if request.session.get('user_role') != 'admin':
        messages.error(request, "Доступ запрещен. Требуются права администратора.")
        return redirect('index')
    # Получение всех мероприятий
    events = PublicEvent.objects.all().order_by('start_date')
    # Рендеринг шаблона панели управления
    return render(request, 'control_panel.html', {'events': events})

# Представление для регистрации ментора (только POST)
@require_http_methods(["POST"])
def register_mentor(request):
    # Создание формы с данными из POST-запроса
    form = MentorRegistrationForm(request.POST)
    if form.is_valid():
        email = form.cleaned_data["email"]
        event_id = form.cleaned_data["event_id"]

        try:
            with connection.cursor() as cursor:
                # Поиск ментора по email
                cursor.execute("SELECT id_mentor FROM mentors WHERE email = %s", [email])
                row = cursor.fetchone()

                if not row:
                    # Создание нового ментора, если не найден
                    cursor.execute("SELECT MAX(id_mentor) FROM mentors")
                    max_id = cursor.fetchone()[0] or 0
                    mentor_id = max_id + 1

                    cursor.execute("""
                        INSERT INTO mentors (id_mentor, surname, name, patronymic, email, phone)
                        VALUES (%s, %s, %s, %s, %s, %s)
                    """, [
                        mentor_id,
                        form.cleaned_data['surname'],
                        form.cleaned_data['name'],
                        form.cleaned_data['patronymic'] or '',
                        email,
                        form.cleaned_data['phone']
                    ])
                else:
                    mentor_id = row[0]

                # Проверка наличия ментора на мероприятии
                cursor.execute("""
                    SELECT id_mentor FROM events
                    WHERE id_event = %s
                """, [event_id])
                existing_mentor = cursor.fetchone()
                if existing_mentor and existing_mentor[0] is not None:
                    messages.error(request, "На данное мероприятие уже зарегистрирован наставник.")
                    request.session['show_mentor_modal'] = True
                    return redirect('index')

                # Привязка ментора к мероприятию
                cursor.execute("""
                    UPDATE events
                    SET id_mentor = %s
                    WHERE id_event = %s AND start_date > CURRENT_DATE
                """, [mentor_id, event_id])

            messages.success(request, "Ментор успешно зарегистрирован на мероприятие.")
        except Exception as e:
            logger.error(f"Ошибка регистрации ментора: {str(e)}")
            messages.error(request, f"Ошибка регистрации ментора: {str(e)}")
            request.session['show_mentor_modal'] = True
    else:
        logger.debug(f"Form errors: {form.errors}")
        messages.error(request, "Пожалуйста, исправьте ошибки в форме.")
        request.session['show_mentor_modal'] = True

    return redirect('index')

# Повторная инициализация логгера
logger = logging.getLogger(__name__)

# Представление для получения списка таблиц и представлений
@require_GET
def get_tables_and_views(request):
    try:
        with connections['org_db'].cursor() as cursor:
            # Получение списка таблиц
            cursor.execute("""
                           SELECT table_name
                           FROM information_schema.tables
                           WHERE table_schema = 'public'
                             AND table_type = 'BASE TABLE'
                           """)
            tables = [row[0] for row in cursor.fetchall()]

            # Получение списка представлений
            cursor.execute("""
                           SELECT table_name
                           FROM information_schema.views
                           WHERE table_schema = 'public'
                           """)
            views = [row[0] for row in cursor.fetchall()]

            # Объединение и сортировка
            all_objects = sorted(tables + views)
            logger.debug(f"Retrieved tables and views: {all_objects}")
            return JsonResponse(all_objects, safe=False, status=200)
    except DatabaseError as e:
        logger.error(f"Error fetching tables and views: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

# Представление для получения списка таблиц
@require_GET
def get_tables(request):
    try:
        with connections['org_db'].cursor() as cursor:
            # Получение списка таблиц
            cursor.execute("""
                           SELECT table_name
                           FROM information_schema.tables
                           WHERE table_schema = 'public'
                             AND table_type = 'BASE TABLE'
                           """)
            tables = [row[0] for row in cursor.fetchall()]
            logger.debug(f"Retrieved tables: {tables}")
            return JsonResponse(tables, safe=False, status=200)
    except DatabaseError as e:
        logger.error(f"Error fetching tables: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

# Представление для получения полей таблицы
@require_GET
def get_table_fields(request):
    # Получение имени таблицы из GET-параметров
    table_name = request.GET.get('table')
    logger.debug(f"Fetching fields for table: {table_name}")

    if not table_name:
        logger.warning("Table name is missing in request.GET")
        return JsonResponse({'success': False, 'error': 'Не указана таблица'}, status=400)

    try:
        with connections['org_db'].cursor() as cursor:
            # Получение списка столбцов таблицы
            cursor.execute("""
                           SELECT column_name
                           FROM information_schema.columns
                           WHERE table_schema = 'public'
                             AND table_name = %s
                           """, [table_name])
            fields = [row[0] for row in cursor.fetchall()]
            logger.debug(f"Fields for {table_name}: {fields}")
            if not fields:
                logger.warning(f"No fields found for table: {table_name}")
                return JsonResponse({'success': False, 'error': 'Таблица или представление не найдены'}, status=404)
            return JsonResponse(fields, safe=False, status=200)
    except DatabaseError as e:
        logger.error(f"Error fetching fields for {table_name}: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

# Представление для получения данных таблицы
@require_GET
def get_table_data(request):
    # Получение имени таблицы из GET-параметров
    table_name = request.GET.get('table')
    logger.debug(f"Fetching data for table: {table_name}")

    if not table_name:
        logger.warning("Table name is missing in request.GET")
        return JsonResponse({'success': False, 'error': 'Не указана таблица'}, status=400)

    try:
        with connections['org_db'].cursor() as cursor:
            # Получение списка столбцов
            cursor.execute("""
                           SELECT column_name
                           FROM information_schema.columns
                           WHERE table_schema = 'public'
                             AND table_name = %s
                           """, [table_name])
            columns = [row[0] for row in cursor.fetchall()]
            if not columns:
                logger.warning(f"No columns found for table: {table_name}")
                return JsonResponse({'success': False, 'error': 'Таблица не найдена'}, status=404)

            # Определение первичного ключа
            pk_field = next((col for col in columns if col.startswith('id_')), columns[0])
            query = f"""
                SELECT {', '.join(columns)}
                FROM public.{table_name}
            """
            cursor.execute(query)
            rows = cursor.fetchall()
            data = [dict(zip(columns, row)) for row in rows]
            for row in data:
                row['id'] = row.pop(pk_field)
            logger.debug(f"Data for {table_name}: {len(data)} rows")
            return JsonResponse({'success': True, 'data': data}, status=200)
    except DatabaseError as e:
        logger.error(f"Error fetching data for {table_name}: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

# Представление для сохранения изменений в таблице
@require_POST
def save_table_data(request):
    try:
        # Парсинг JSON-данных из тела запроса
        data = json.loads(request.body)
        table_name = data.get('table')
        record_id = data.get('id')
        fields = data.get('fields')
        logger.debug(f"Saving data: table={table_name}, id={record_id}, fields={fields}")

        if not table_name or not record_id or not fields:
            logger.warning("Missing required fields in save_table_data")
            return JsonResponse({'success': False, 'error': 'Не указаны обязательные поля'}, status=400)

        try:
            with connections['org_db'].cursor() as cursor:
                # Получение списка столбцов таблицы
                cursor.execute("""
                               SELECT column_name
                               FROM information_schema.columns
                               WHERE table_schema = 'public'
                                 AND table_name = %s
                               """, [table_name])
                valid_columns = [row[0] for row in cursor.fetchall()]
                if not valid_columns:
                    logger.warning(f"No columns found for table: {table_name}")
                    return JsonResponse({'success': False, 'error': 'Таблица не найдена'}, status=404)

                # Определение первичного ключа
                pk_field = next((col for col in valid_columns if col.startswith('id_')), valid_columns[0])
                valid_fields = {k: v for k, v in fields.items() if k in valid_columns and k != pk_field}
                if not valid_fields:
                    logger.warning("No valid fields provided for update")
                    return JsonResponse({'success': False, 'error': 'Нет допустимых полей для обновления'}, status=400)

                # Формирование SQL-запроса для обновления
                set_clause = ', '.join(f"{k} = %s" for k in valid_fields.keys())
                query = f"""
                    UPDATE public.{table_name}
                    SET {set_clause}
                    WHERE {pk_field} = %s
                """
                params = list(valid_fields.values()) + [record_id]
                cursor.execute(query, params)
                if cursor.rowcount == 0:
                    logger.warning(f"No rows updated for table={table_name}, id={record_id}")
                    return JsonResponse({'success': False, 'error': 'Запись не найдена'}, status=404)
            return JsonResponse({'success': True}, status=200)
        except DatabaseError as e:
            logger.error(f"Error saving data for {table_name}: {str(e)}")
            return JsonResponse({'success': False, 'error': str(e)}, status=500)
    except json.JSONDecodeError as e:
        logger.error(f"Invalid JSON in save_table_data: {str(e)}")
        return JsonResponse({'success': False, 'error': 'Неверный формат данных'}, status=400)

# Представление для добавления новой записи в таблицу
@require_POST
def add_record(request):
    try:
        # Получение данных из POST-запроса
        data = request.POST
        table_name = data.get('table_name')
        fields = {k: v for k, v in data.items() if
                  k not in ['csrfmiddlewaretoken', 'record_submit', 'table_name', 'record_id']}
        logger.debug(f"Adding record: table={table_name}, fields={fields}")

        if not table_name or not fields:
            logger.warning("Missing required fields in add_record")
            return JsonResponse({'success': False, 'error': 'Не указаны обязательные поля'}, status=400)

        try:
            with connections['org_db'].cursor() as cursor:
                # Получение списка столбцов таблицы
                cursor.execute("""
                               SELECT column_name
                               FROM information_schema.columns
                               WHERE table_schema = 'public'
                                 AND table_name = %s
                               """, [table_name])
                valid_columns = [row[0] for row in cursor.fetchall()]
                if not valid_columns:
                    logger.warning(f"No columns found for table: {table_name}")
                    return JsonResponse({'success': False, 'error': 'Таблица не найдена'}, status=404)

                # Определение первичного ключа
                pk_field = next((col for col in valid_columns if col.startswith('id_')), valid_columns[0])
                valid_fields = {k: v for k, v in fields.items() if k in valid_columns and k != pk_field}
                if not valid_fields:
                    logger.warning("No valid fields provided for insert")
                    return JsonResponse({'success': False, 'error': 'Нет допустимых полей для добавления'}, status=400)

                # Генерация нового ID
                cursor.execute(f"SELECT MAX({pk_field}) FROM public.{table_name}")
                max_id = cursor.fetchone()[0] or 0
                new_id = max_id + 1

                # Формирование SQL-запроса для вставки
                columns = ', '.join([pk_field] + list(valid_fields.keys()))
                placeholders = ', '.join(['%s'] * (len(valid_fields) + 1))
                query = f"""
                    INSERT INTO public.{table_name} ({columns})
                    VALUES ({placeholders})
                """
                params = [new_id] + list(valid_fields.values())
                cursor.execute(query, params)

            return JsonResponse({'success': True}, status=200)
        except DatabaseError as e:
            logger.error(f"Error adding record to {table_name}: {str(e)}")
            return JsonResponse({'success': False, 'error': str(e)}, status=500)
    except Exception as e:
        logger.error(f"Error in add_record: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)}, status=400)

# Повторная инициализация логгера
logger = logging.getLogger(__name__)

# Представление для генерации отчетов
@require_POST
def generate_summary(request):
    # Получение данных из POST-запроса
    table_name = request.POST.get('table')
    action = request.POST.get('action')
    file_format = request.POST.get('format')

    logger.debug(f"Received POST data: table={table_name}, action={action}, format={file_format}")

    # Проверка наличия имени таблицы
    if not table_name:
        logger.warning("Table name is missing in request.POST")
        return JsonResponse({'success': False, 'error': 'Не указано представление'}, status=400)

    # Проверка допустимого действия
    if action not in ['create', 'display']:
        logger.warning(f"Invalid action: {action}")
        return JsonResponse({'success': False, 'error': 'Недопустимое действие'}, status=400)

    # Проверка формата файла для действия create
    if action == 'create' and file_format not in ['pdf', 'word']:
        logger.warning(f"Invalid format: {file_format}")
        return JsonResponse({'success': False, 'error': 'Недопустимый формат отчета'}, status=400)

    # Проверка существования представления
    try:
        with connections['org_db'].cursor() as cursor:
            cursor.execute("""
                SELECT 1
                FROM information_schema.views
                WHERE table_schema = 'public' AND table_name = %s
            """, [table_name])
            if not cursor.fetchone():
                logger.warning(f"View does not exist: {table_name}")
                return JsonResponse({'success': False, 'error': 'Представление не найдено'}, status=404)
    except DatabaseError as e:
        logger.error(f"Error checking view existence: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

    # Получение столбцов представления
    try:
        with connections['org_db'].cursor() as cursor:
            cursor.execute("""
                SELECT column_name
                FROM information_schema.columns
                WHERE table_schema = 'public' AND table_name = %s
            """, [table_name])
            columns = [row[0] for row in cursor.fetchall()]
            if not columns:
                logger.warning(f"No columns found for {table_name}")
                return JsonResponse({'success': False, 'error': 'Столбцы не найдены'}, status=404)
    except DatabaseError as e:
        logger.error(f"Error fetching columns for {table_name}: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

    # Формирование SQL-запроса
    query = f"SELECT {', '.join(columns)} FROM public.{table_name}"
    params = []

    try:
        with connections['org_db'].cursor() as cursor:
            logger.debug(f"Executing query: {query}")
            cursor.execute(query, params)
            rows = cursor.fetchall()
            data = [dict(zip(columns, row)) for row in rows]
            logger.debug(f"Retrieved {len(data)} rows for {table_name}")

            if action == 'display':
                # Возврат данных в формате JSON для отображения
                return JsonResponse({'success': True, 'data': data}, status=200)

            # Генерация имени файла с временной меткой
            timestamp = int(time.time())
            file_name = f"{table_name}_report_{timestamp}.pdf" if file_format == 'pdf' else f"{table_name}_report_{timestamp}.docx"

            if file_format == 'pdf':
                try:
                    # Создание PDF-документа
                    buffer = io.BytesIO()
                    doc = SimpleDocTemplate(buffer, pagesize=letter)
                    elements = []

                    # Регистрация шрифта DejaVuSans
                    font_name = 'DejaVuSans'
                    font_path = os.path.join(settings.BASE_DIR, 'static', 'fonts', 'DejaVuSans.ttf')
                    if os.path.exists(font_path):
                        pdfmetrics.registerFont(TTFont(font_name, font_path))
                        logger.debug(f"Registered font {font_name} at {font_path}")
                    else:
                        font_name = 'Helvetica'
                        logger.warning(f"Font file not found at {font_path}, falling back to Helvetica")

                    # Настройка стиля заголовка
                    styles = getSampleStyleSheet()
                    title_style = styles['Title']
                    title_style.fontName = font_name
                    title_style.fontSize = 16

                    elements.append(Paragraph(f"Отчет по {table_name}", title_style))

                    # Формирование таблицы с данными
                    table_data = [columns]
                    for row in data:
                        table_data.append([str(row[col]) if row[col] is not None else '' for col in columns])

                    table = Table(table_data)
                    table.setStyle(TableStyle([
                        ('FONTNAME', (0, 0), (-1, -1), font_name),
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTSIZE', (0, 0), (-1, 0), 12),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ]))
                    elements.append(table)
                    doc.build(elements)

                    buffer.seek(0)
                    # Возврат PDF-файла
                    return FileResponse(
                        buffer,
                        as_attachment=True,
                        filename=file_name,
                        content_type='application/pdf'
                    )
                except Exception as e:
                    logger.error(f"Error generating PDF for {table_name}: {str(e)}")
                    return JsonResponse({'success': False, 'error': f"Ошибка генерации PDF: {str(e)}"}, status=400)

            elif file_format == 'word':
                try:
                    # Создание Word-документа
                    doc = docx.Document()
                    doc.add_heading(f'Отчет по {table_name}', 0)

                    # Создание таблицы
                    table = doc.add_table(rows=len(data) + 1, cols=len(columns))
                    table.style = 'Table Grid'

                    # Заполнение заголовков таблицы
                    for i, col in enumerate(columns):
                        table.cell(0, i).text = col

                    # Заполнение данных таблицы
                    for i, row in enumerate(data, 1):
                        for j, col in enumerate(columns):
                            table.cell(i, j).text = str(row[col] if row[col] is not None else '')

                    buffer = io.BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    # Возврат Word-файла
                    return FileResponse(
                        buffer,
                        as_attachment=True,
                        filename=file_name,
                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                except Exception as e:
                    logger.error(f"Error generating Word for {file_name}: {str(e)}")
                    return JsonResponse({'success': False, 'error': f"Ошибка генерации Word: {str(e)}"}, status=400)

    except DatabaseError as e:
        logger.error(f"Database error generating report for {table_name}: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)
    except Exception as e:
        logger.error(f"Unexpected error generating report for {table_name}: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

# Представление для получения списка представлений
@require_GET
def get_views(request):
    try:
        with connections['org_db'].cursor() as cursor:
            # Получение списка представлений
            cursor.execute("""
                SELECT table_name 
                FROM information_schema.views 
                WHERE table_schema = 'public'
            """)
            views = [row[0] for row in cursor.fetchall()]
            logger.debug(f"Retrieved views: {views}")
            return JsonResponse(views, safe=False, status=200)
    except DatabaseError as e:
        logger.error(f"Error fetching views: {str(e)}")
        return JsonResponse({'success': False, 'error': str(e)}, status=500)